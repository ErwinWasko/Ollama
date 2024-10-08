from flask import Flask, render_template, jsonify, request, send_file
import skaner
import ollama
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import io
from docx import Document
import mysql.connector
from mysql.connector import Error
from reportlab.lib.pagesizes import letter

app = Flask(__name__)

stop_generation_flag = False
generation_thread = None

client = ollama.Client()
model_name = "llama3"

# Przechowuje bieżące raporty, które są analizowane
current_reports = []
all_reports_fetched = False

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/ask_chatbot', methods=['POST'])
def ask_chatbot():
    try:
        # Pobierz pytanie z frontendu
        data = request.json
        user_input = data.get('question')

        if not user_input:
            return jsonify({'error': 'Nie podano pytania'}), 400

        # Przygotowanie klienta Ollama
        client = ollama.Client()
        model_name = "llama3"

        # Wywołanie modelu z wpisanym przez użytkownika pytaniem
        response = client.generate(model=model_name, prompt=user_input)

        if isinstance(response, dict) and 'response' in response:
            return jsonify({'answer': response['response']})
        else:
            return jsonify({'error': 'Nie otrzymano odpowiedzi od modelu'}), 500
    
    except Exception as e:
        # Debugowanie w razie błędu
        print(f"Błąd podczas przetwarzania zapytania: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/fetch_reports')
def fetch_reports():
    global current_reports, all_reports_fetched, stop_generation_flag

    # Pobranie parametru cvss z żądania
    selected_cvss = request.args.get('cvss', 'all')

    # Inicjalizacja raportów, jeśli są puste i jeszcze nie pobraliśmy wszystkich
    if not current_reports and not all_reports_fetched:
        current_reports = skaner.fetch_security_reports()

    total_reports = len(current_reports)  # Liczba raportów w bazie
    max_cvss = max([float(report['vulnerability_score']) for report in current_reports]) if current_reports else 0  # Maksymalna podatność

    # Sprawdzenie, czy wszystkie raporty zostały przetworzone
    if all_reports_fetched or not current_reports or not current_reports:
        return jsonify({
            'done': True,
            'total_reports': total_reports,
            'max_cvss': max_cvss
        })  # Zwróć informację o zakończeniu

    # Przetwarzanie raportów jeden po drugim
    report = current_reports.pop(0)
    cve = report.get('vulnerability')
    cvss = float(report.get('vulnerability_score', 0))
    description = report.get('vulnerability_description', 'No description available')

    # Pobieranie danych z MITRE
    mitre_data = skaner.fetch_mitre_data(cve)

    # Przetwarzanie danych z MITRE, aby uzyskać referencje
    description, ref_urls = skaner.process_mitre_data(mitre_data)

    # Przetwarzanie analizy Ollamy z referencjami
    ollama_analysis = skaner.analyze_data_with_ollama(cve, cvss, description, ref_urls)

    result = {
        'cve': cve,
        'cvss': cvss,
        'description': description,
        'ollama_analysis': ollama_analysis,
        'cve_link': f'https://www.cve.org/CVERecord?id={cve}'
    }

    if not current_reports:
        all_reports_fetched = True

    return jsonify({
        'done': False,
        'result': result,
        'total_reports': total_reports,
        'max_cvss': max_cvss
    })

@app.route('/stop_generating', methods=['POST'])
def stop_generating():
    global stop_generation_flag
    stop_generation_flag = True
    return jsonify({"message": "Report generation stopped."})

# Funkcja do generowania pliku PDF
@app.route('/generate_pdf_report', methods=['POST'])
def generate_pdf_report():
    data = request.json
    reports = data.get('reports', [])

    # Tworzenie bufora dla pliku PDF
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y_position = height - 50  # Startowa pozycja Y

    # Funkcja do zawijania długich linii tekstu
    def draw_wrapped_text(pdf, text, x, y, max_width):
        lines = []
        words = text.split(' ')
        current_line = []
        current_width = 0

        for word in words:
            word_width = pdf.stringWidth(word + ' ', 'Helvetica', 12)
            if current_width + word_width > max_width:
                lines.append(' '.join(current_line))
                current_line = [word]
                current_width = word_width
            else:
                current_line.append(word)
                current_width += word_width

        if current_line:
            lines.append(' '.join(current_line))

        for line in lines:
            pdf.drawString(x, y, line)
            y -= 15  # Przesunięcie w dół dla kolejnej linii

        return y  # Zwróć zaktualizowaną pozycję Y

    # Funkcja do czyszczenia tekstu z nieprawidłowych znaków
    def clean_text(text):
        if not text:
            return ''
        return ''.join(char for char in text if char.isprintable())  # Usuwa nieczytelne znaki

    # Ustawienia marginesów
    left_margin = 50
    right_margin = width - 50
    max_text_width = right_margin - left_margin

    # Rysowanie każdego raportu
    for report in reports:
        if y_position < 100:  # Jeśli miejsce na stronie się kończy, twórz nową stronę
            pdf.showPage()
            y_position = height - 50

        pdf.setFont("Helvetica-Bold", 12)
        pdf.drawString(left_margin, y_position, clean_text(f"CVE: {report['cve']}"))
        y_position -= 20
        
        pdf.setFont("Helvetica", 10)
        y_position = draw_wrapped_text(pdf, clean_text(f"CVSS Score: {report['cvss']}"), left_margin, y_position, max_text_width)
        y_position = draw_wrapped_text(pdf, clean_text(f"Description: {report['description']}"), left_margin, y_position, max_text_width)
        y_position = draw_wrapped_text(pdf, clean_text(f"Ollama Analysis: {report['ollama_analysis']}"), left_margin, y_position, max_text_width)
        
        y_position -= 40  # Przesunięcie dla kolejnego raportu

    pdf.save()

    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="Raporty_CVSS.pdf", mimetype='application/pdf')

# Funkcja do generowania pliku Word
@app.route('/generate_word_report', methods=['POST'])
def generate_word_report():
    data = request.json
    reports = data.get('reports', [])
    
    document = Document()
    document.add_heading('Raporty CVSS', 0)

    # Dodawanie każdego raportu do pliku Word
    for report in reports:
        document.add_heading(f'Report for {report["cve"]}', level=1)
        document.add_paragraph(f'CVSS Score: {report["cvss"]}')
        document.add_paragraph(f'Description: {report["description"]}')
        document.add_paragraph(f'Ollama Analysis: {report["ollama_analysis"]}')
        document.add_paragraph('')

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="Raporty_CVSS.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@app.route('/fetch_chart_data')
def fetch_chart_data():
    # Połącz się z bazą danych i pobierz raporty
    reports = skaner.fetch_security_reports()

    if not reports:
        return jsonify({'error': 'No data available'})

    # Grupowanie raportów według zakresów CVSS
    cvss_ranges = {
        'Low': 0,
        'Medium': 0,
        'High': 0,
        'Critical': 0
    }

    for report in reports:
        score = float(report.get('vulnerability_score', 0))
        if 0.1 <= score <= 3.9:
            cvss_ranges['Low'] += 1
        elif 4.0 <= score <= 6.9:
            cvss_ranges['Medium'] += 1
        elif 7.0 <= score <= 8.9:
            cvss_ranges['High'] += 1
        elif 9.0 <= score <= 10.0:
            cvss_ranges['Critical'] += 1

    return jsonify({
        'cvss_ranges': cvss_ranges
    })

# Funkcja łącząca się z bazą danych
def connect_to_database():
    try:
        connection = mysql.connector.connect(
            host='localhost',
            database='test',
            user='root',
            password=''
        )
        return connection
    except Error as e:
        print(f"Error: {e}")
        return None

# Funkcja pobierająca listę podatności (CVE) z bazy danych
def fetch_cve_list():
    connection = connect_to_database()
    if connection is None:
        return []

    try:
        cursor = connection.cursor()
        query = "SELECT vulnerability FROM vulnerabilities"
        cursor.execute(query)
        results = [row[0] for row in cursor.fetchall()]
        return results
    except Error as e:
        print(f"Błąd podczas pobierania CVE: {e}")
        return []
    finally:
        if connection:
            cursor.close()
            connection.close()

@app.route('/get_cve_list', methods=['GET'])
def get_cve_list():
    # Pobierz listę CVE z bazy danych
    cve_list = fetch_cve_list()
    
    if not cve_list:
        return jsonify({'error': 'Brak podatności w bazie danych'}), 500
    
    return jsonify({'cve_list': cve_list})

@app.route('/simulate_attack', methods=['POST'])
def simulate_attack():
    data = request.json
    vulnerability = data.get('vulnerability')

    if not vulnerability:
        return jsonify({'error': 'Brak podatności'}), 400

    # Pobieramy szczegóły podatności z bazy danych
    vulnerability_details = skaner.fetch_vulnerability_details(vulnerability)

    if not vulnerability_details:
        return jsonify({'error': 'Nie znaleziono szczegółów dla podanego CVE'}), 404

    # Rozpakowanie wyników z bazy danych
    cvss_score, description = vulnerability_details  

    # Generowanie scenariusza ataku z dodatkowymi danymi
    attack_scenario = skaner.generate_attack_scenario(vulnerability, cvss_score, description)
    
    return jsonify({'steps': attack_scenario})

if __name__ == '__main__':
    app.run(debug=True)
