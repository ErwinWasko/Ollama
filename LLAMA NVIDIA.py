SKANER.PY:

import mysql.connector
from mysql.connector import Error
import requests
from openai import OpenAI

# Funkcja łącząca się z bazą danych
def connect_to_database():
    try:
        connection = mysql.connector.connect(
            host='localhost',
            database='test',
            user='root',
            password=''
        )
        return connection
    except Error as e:
        print(f"Error: {e}")
        return None

# Funkcja pobierająca raporty o podatnościach z bazy danych
def fetch_security_reports():
    connection = connect_to_database()
    if connection is None:
        return [], 0  # Zwraca pustą listę i 0, jeśli połączenie nie działa

    try:
        cursor = connection.cursor()
        query = """
        SELECT * 
        FROM vulnerabilities 
        ORDER BY CAST(vulnerability_score AS DECIMAL(3,1)) DESC
        """
        cursor.execute(query)
        columns = [desc[0] for desc in cursor.description]
        results = [dict(zip(columns, row)) for row in cursor.fetchall()]
        return results
    except Error as e:
        print(f"Error: {e}")
        return [], 0
    finally:
        if connection:
            cursor.close()
            connection.close()

def fetch_vulnerability_details(cve):
    connection = connect_to_database()
    if connection is None:
        return None

    try:
        cursor = connection.cursor()
        query = """
        SELECT vulnerability_score, vulnerability_description
        FROM vulnerabilities 
        WHERE vulnerability = %s
        """
        cursor.execute(query, (cve,))
        result = cursor.fetchone()

        if result:
            cvss_score = result[0]
            description = result[1]
            return cvss_score, description  
        else:
            return None

    except Error as e:
        print(f"Error: {e}")
        return None
    finally:
        if connection:
            cursor.close()
            connection.close()

def fetch_mitre_data(cve):
    url = f'https://cveawg.mitre.org/api/cve/{cve}'
    try:
        response = requests.get(url)
        response.raise_for_status()
        mitre_data = response.json()
        
        if mitre_data and 'cveMetadata' in mitre_data:
            return mitre_data
        else:
            return None
    except requests.RequestException as e:
        print(f"Error fetching data from MITRE: {e}")
        return None
    except Exception as e:
        print(f"An error occurred: {e}")
        return None

def process_mitre_data(mitre_data):
    if not mitre_data:
        return "No data available", []

    description = mitre_data.get('containers', {}).get('cna', {}).get('descriptions', [{'value': 'No description available'}])[0]['value']
    references = mitre_data.get('containers', {}).get('cna', {}).get('references', [])
    
    ref_urls = [ref['url'] for ref in references]
    
    return description, ref_urls

def analyze_data_with_nvidia_llama(cve, cvss, description, ref_urls):
    try:
        # Przygotowanie klienta NVIDIA LLaMA
        client = OpenAI(
            base_url = "https://integrate.api.nvidia.com/v1",
            api_key = "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"
        )

        # Konstruowanie wiadomości
        ref_text = "\n".join([f"- {url}" for url in ref_urls])
        prompt = f"""
        Tutaj znajdują się szczegóły podatności:
        CVE: {cve}
        CVSS Score: {cvss}
        Opis: {description}
        
        Dodatkowo, oto adresy URL referencji do dalszej analizy:
        {ref_text}

        Proszę przeanalizować te informacje i zasugerować konkretne kroki w celu zmniejszenia oceny CVSS oraz złagodzenia tej podatności na podstawie treści z podanych stron. Potwierdź, że analiza została przeprowadzona na podstawie dostarczonych URL i podsumuj wyciągnięte wnioski.
        """

        # Wysłanie zapytania do modelu NVIDIA LLaMA
        completion = client.chat.completions.create(
            model="meta/llama-3.1-405b-instruct",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            top_p=0.7,
            max_tokens=1024,
            stream=False
        )

        # Pobranie odpowiedzi
        response = completion.choices[0].message.content.strip()

        return f"NVIDIA LLaMA przeanalizowała dostarczone URL i wyciągnęła wnioski:\n{response}"

    except Exception as e:
        print(f"Error analyzing data with NVIDIA LLaMA: {e}")
        return "No valid response from NVIDIA LLaMA."

def generate_attack_scenario(cve, cvss_score, description):
    try:
        # Przygotowanie klienta NVIDIA LLaMA
        client = OpenAI(
            base_url = "https://integrate.api.nvidia.com/v1",
            api_key = "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"
        )

        # Konstrukcja promptu dla modelu
        prompt = f"""
        Wygeneruj krok po kroku scenariusz ataku dla następującej podatności:
        CVE: {cve}
        CVSS Score: {cvss_score}
        Opis: {description}

        Uwzględnij kroki odkrycia, eksploatacji, eskalacji uprawnień i wpływu.
        """

        # Wysłanie zapytania do modelu NVIDIA LLaMA
        completion = client.chat.completions.create(
            model="meta/llama-3.1-405b-instruct",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            top_p=0.7,
            max_tokens=1024,
            stream=False
        )

        # Pobranie odpowiedzi i przetworzenie na listę kroków
        response = completion.choices[0].message.content.strip()
        steps = response.split('\n')
        return steps

    except Exception as e:
        print(f"Error generating attack scenario: {e}")
        return ["Error generating attack scenario."]

def main():
    # Pobieranie raportów o podatnościach
    reports = fetch_security_reports()
    
    if not isinstance(reports, list) or not reports:
        print("No security reports found.")
        return
    
    # Iterowanie przez wszystkie raporty
    for report in reports:
        cve = report.get('vulnerability')
        cvss = report.get('vulnerability_score', 'N/A')
        description = report.get('vulnerability_description', 'No description available')
        
        # Pobieranie danych z MITRE CVE API
        mitre_data = fetch_mitre_data(cve)
        description, ref_urls = process_mitre_data(mitre_data)
        
        # Analiza danych przy użyciu NVIDIA LLaMA
        nvidia_analysis = analyze_data_with_nvidia_llama(cve, cvss, description, ref_urls)
        
        # Generowanie scenariusza ataku
        attack_scenario = generate_attack_scenario(cve, cvss, description)
        
        # Wyświetl wyniki
        print(f"Raport dla CVE: {cve}")
        print("Analiza NVIDIA LLaMA:")
        print(nvidia_analysis if nvidia_analysis else "No analysis available")
        print("\nScenariusz ataku:")
        for step in attack_scenario:
            print(step)
        print("\n")

if __name__ == "__main__":
    main()


APP.PY:

from flask import Flask, render_template, jsonify, request, send_file
import skaner
import ollama
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import io
from docx import Document
import mysql.connector
from mysql.connector import Error

app = Flask(__name__)

client = ollama.Client()
model_name = "llama3"

# Przechowuje bieżące raporty, które są analizowane
current_reports = []
all_reports_fetched = False

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/ask_chatbot', methods=['POST'])
def ask_chatbot():
    try:
        # Pobierz pytanie z frontendu
        data = request.json
        user_input = data.get('question')
        
        # Debug: wyświetlenie pytania, aby upewnić się, że dane docierają do serwera
        print(f"Received question: {user_input}")

        if not user_input:
            return jsonify({'error': 'Nie podano pytania'}), 400

        # Wywołanie funkcji z pliku skaner.py, która analizuje pytanie
        response = skaner.analyze_data_with_nvidia_llama(user_input, 0, '', [])

        # Debug: wyświetlenie odpowiedzi, aby upewnić się, że Ollama działa poprawnie
        print(f"Ollama response: {response}")

        return jsonify({'answer': response})
    
    except Exception as e:
        # Debugowanie w razie błędu
        print(f"Błąd podczas przetwarzania zapytania: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/fetch_reports')
def fetch_reports():
    global current_reports, all_reports_fetched

    # Pobranie parametru cvss z żądania
    selected_cvss = request.args.get('cvss', 'all')

    # Inicjalizacja raportów, jeśli są puste i jeszcze nie pobraliśmy wszystkich
    if not current_reports and not all_reports_fetched:
        current_reports = skaner.fetch_security_reports()

    total_reports = len(current_reports)  # Liczba raportów w bazie
    max_cvss = max([float(report['vulnerability_score']) for report in current_reports]) if current_reports else 0  # Maksymalna podatność

    # Sprawdzenie, czy wszystkie raporty zostały przetworzone
    if all_reports_fetched or not current_reports:
        return jsonify({
            'done': True,
            'total_reports': total_reports,
            'max_cvss': max_cvss
        })  # Zwróć informację o zakończeniu

    # Przetwarzanie raportów jeden po drugim
    report = current_reports.pop(0)
    cve = report.get('vulnerability')
    cvss = float(report.get('vulnerability_score', 0))
    description = report.get('vulnerability_description', 'No description available')

    # Przetwarzanie analizy Ollamy
    ollama_analysis = skaner.analyze_data_with_nvidia_llama(cve, cvss, description, [])

    result = {
        'cve': cve,
        'cvss': cvss,
        'description': description,
        'ollama_analysis': ollama_analysis,
        'cve_link': f'https://www.cve.org/CVERecord?id={cve}'
    }

    if not current_reports:
        all_reports_fetched = True

    return jsonify({
        'done': False,
        'result': result,
        'total_reports': total_reports,
        'max_cvss': max_cvss
    })

# Funkcja do generowania pliku PDF
@app.route('/generate_pdf_report', methods=['POST'])
def generate_pdf_report():
    data = request.json
    reports = data.get('reports', [])
    
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y_position = height - 50

    # Rysowanie każdego raportu
    for report in reports:
        c.drawString(100, y_position, f"CVE: {report['cve']}")
        c.drawString(100, y_position - 20, f"CVSS Score: {report['cvss']}")
        c.drawString(100, y_position - 40, f"Description: {report['description']}")
        c.drawString(100, y_position - 60, f"Ollama Analysis: {report['ollama_analysis']}")
        y_position -= 100  # Przesunięcie w dół dla kolejnego raportu
        if y_position < 100:
            c.showPage()
            y_position = height - 50

    c.save()

    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="Raporty_CVSS.pdf", mimetype='application/pdf')

# Funkcja do generowania pliku Word
@app.route('/generate_word_report', methods=['POST'])
def generate_word_report():
    data = request.json
    reports = data.get('reports', [])
    
    document = Document()
    document.add_heading('Raporty CVSS', 0)

    # Dodawanie każdego raportu do pliku Word
    for report in reports:
        document.add_heading(f'Report for {report["cve"]}', level=1)
        document.add_paragraph(f'CVSS Score: {report["cvss"]}')
        document.add_paragraph(f'Description: {report["description"]}')
        document.add_paragraph(f'Ollama Analysis: {report["ollama_analysis"]}')
        document.add_paragraph('')

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="Raporty_CVSS.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@app.route('/fetch_chart_data')
def fetch_chart_data():
    # Połącz się z bazą danych i pobierz raporty
    reports = skaner.fetch_security_reports()

    if not reports:
        return jsonify({'error': 'No data available'})

    # Grupowanie raportów według zakresów CVSS
    cvss_ranges = {
        'Low': 0,
        'Medium': 0,
        'High': 0,
        'Critical': 0
    }

    for report in reports:
        score = float(report.get('vulnerability_score', 0))
        if 0.1 <= score <= 3.9:
            cvss_ranges['Low'] += 1
        elif 4.0 <= score <= 6.9:
            cvss_ranges['Medium'] += 1
        elif 7.0 <= score <= 8.9:
            cvss_ranges['High'] += 1
        elif 9.0 <= score <= 10.0:
            cvss_ranges['Critical'] += 1

    return jsonify({
        'cvss_ranges': cvss_ranges
    })

# Funkcja łącząca się z bazą danych
def connect_to_database():
    try:
        connection = mysql.connector.connect(
            host='localhost',
            database='test',
            user='root',
            password=''
        )
        return connection
    except Error as e:
        print(f"Error: {e}")
        return None

# Funkcja pobierająca listę podatności (CVE) z bazy danych
def fetch_cve_list():
    connection = connect_to_database()
    if connection is None:
        return []

    try:
        cursor = connection.cursor()
        query = "SELECT vulnerability FROM vulnerabilities"
        cursor.execute(query)
        results = [row[0] for row in cursor.fetchall()]
        return results
    except Error as e:
        print(f"Błąd podczas pobierania CVE: {e}")
        return []
    finally:
        if connection:
            cursor.close()
            connection.close()

@app.route('/get_cve_list', methods=['GET'])
def get_cve_list():
    # Pobierz listę CVE z bazy danych
    cve_list = fetch_cve_list()
    
    if not cve_list:
        return jsonify({'error': 'Brak podatności w bazie danych'}), 500
    
    return jsonify({'cve_list': cve_list})

@app.route('/simulate_attack', methods=['POST'])
def simulate_attack():
    data = request.json
    vulnerability = data.get('vulnerability')

    if not vulnerability:
        return jsonify({'error': 'Brak podatności'}), 400

    # Pobieramy szczegóły podatności z bazy danych
    vulnerability_details = skaner.fetch_vulnerability_details(vulnerability)

    if not vulnerability_details:
        return jsonify({'error': 'Nie znaleziono szczegółów dla podanego CVE'}), 404

    # Rozpakowanie wyników z bazy danych
    cvss_score, description = vulnerability_details  

    # Generowanie scenariusza ataku z dodatkowymi danymi
    attack_scenario = skaner.generate_attack_scenario(vulnerability, cvss_score, description)
    
    return jsonify({'steps': attack_scenario})

if __name__ == '__main__':
    app.run(debug=True)
